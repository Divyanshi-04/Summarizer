import streamlit as st
import torch
from transformers import pipeline
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import docx
import re
import os
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Page configuration
st.set_page_config(
    page_title="Advanced Document Summarizer",
    page_icon="📄",
    layout="wide"
)

# Preprocessing and optimization improvements
def preprocess_text(text):
    """Enhanced text preprocessing with more robust cleaning."""
    # Remove extra whitespaces and normalize text
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove special characters while preserving key punctuation
    text = re.sub(r'[^\w\s.,!?:;()-]', '', text)
    # Ensure proper sentence spacing
    text = re.sub(r'([.!?])([a-zA-Z])', r'\1 \2', text)
    return text

# Load an efficient, free summarization model
@st.cache_resource
def load_summarizer():
    """Load a lightweight, efficient summarization model."""
    try:
        # Use facebook/bart-large-cnn for better performance
        device = "cuda" if torch.cuda.is_available() else "cpu"
        return pipeline("summarization", 
                        model="facebook/bart-large-cnn", 
                        device=0 if device == "cuda" else -1,
                        max_length=1024)  # Increased max length
    except Exception as e:
        st.error(f"Error loading model: {e}")
        return None

# Advanced summarization with intelligent chunking
def advanced_summarize(text, max_length=500, min_length=100):
    """Improved summarization with better chunk handling."""
    if not text or len(text.split()) < 50:
        return "Insufficient text for summarization."
    
    try:
        # Preprocess text
        preprocessed_text = preprocess_text(text)

        # Intelligent chunking with overlap
        chunks = chunk_text(preprocessed_text, chunk_size=2048, overlap=512)
        summaries = []

        with st.spinner("Generating Summary..."):
            for chunk in chunks:
                # Use more advanced summarization parameters
                summary = summarizer(
                    chunk, 
                    max_length=max_length, 
                    min_length=min_length, 
                    do_sample=False,
                    num_beams=4,  # Improved beam search
                    early_stopping=True
                )[0]['summary_text']
                summaries.append(summary)

        # Combine and refine summaries
        final_summary = " ".join(summaries)
        return summarizer(final_summary, 
                          max_length=max_length*2, 
                          min_length=min_length, 
                          do_sample=False)[0]['summary_text']
    
    except Exception as e:
        st.error(f"Summarization error: {e}")
        return "Could not generate summary."

# Enhanced chunking with overlap
def chunk_text(text, chunk_size=2048, overlap=512):
    """Improved text chunking with overlap to preserve context."""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append(chunk)
    
    return chunks

# PDF text extraction with multiple strategies
def extract_text_from_pdf(file):
    """Robust PDF text extraction with fallback mechanisms."""
    try:
        # First, try standard text extraction
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        
        # If text is too sparse, fall back to OCR
        if len(text.split()) < 50:
            st.warning("Low-quality text extraction. Attempting OCR...")
            return ocr_pdf(file)
        
        return text.strip()
    
    except Exception as e:
        st.warning(f"PDF extraction failed. Attempting OCR: {e}")
        return ocr_pdf(file)

def ocr_pdf(file):
    """Optical Character Recognition for PDFs."""
    try:
        # Convert PDF to images with higher resolution
        images = convert_from_path(file, dpi=300)
        
        # Use OCR with improved configuration
        text = ''
        for image in images:
            ocr_text = pytesseract.image_to_string(
                image, 
                config='--oem 3 --psm 6'  # Improved OCR engine mode and page segmentation
            )
            text += ocr_text + "\n"
        
        return text.strip()
    
    except Exception as e:
        st.error(f"OCR failed: {e}")
        return ""

def extract_text_from_docx(file):
    """Improved DOCX text extraction."""
    try:
        doc = docx.Document(file)
        
        # Extract text with more robust method
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Optional: Extract text from tables if needed
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                tables_text.extend([cell.text for cell in row.cells if cell.text.strip()])
        
        return "\n".join(paragraphs + tables_text)
    
    except Exception as e:
        st.error(f"DOCX extraction error: {e}")
        return ""

# Main summarization logic
def generate_structured_summary(text):
    """Create a comprehensive, structured summary."""
    main_summary = advanced_summarize(text, max_length=700, min_length=250)
    
    # Generate key insights using TF-IDF
    sentences = re.split(r'(?<=[.!?])\s+', text)
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(sentences)
    
    # Calculate sentence importance
    sentence_scores = np.sum(tfidf_matrix.toarray(), axis=1)
    top_sentence_indices = sentence_scores.argsort()[-5:][::-1]
    key_sentences = [sentences[i] for i in top_sentence_indices]
    
    # Structured summary template
    structured_summary = f"""# Document Summary

## Overview
{main_summary}

## Key Insights
"""
    for i, sentence in enumerate(key_sentences, 1):
        structured_summary += f"{i}. {sentence.strip()}\n"
    
    return structured_summary

# Streamlit Main App
def main():
    st.title("📄 Advanced Document Summarizer")
    st.write("Intelligent document summarization for large documents")

    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Summary Settings")
        summary_strategy = st.selectbox(
            "Summarization Strategy", 
            ["Concise", "Comprehensive", "Extract Key Points"]
        )
        summary_length = st.slider(
            "Summary Length", 
            100, 1000, 500, 50, 
            help="Adjust the summary's detail level"
        )

    # File uploader with expanded support
    uploaded_file = st.file_uploader(
        "Choose a file", 
        type=["pdf", "docx", "txt"], 
        help="Upload PDF, Word, or Text documents up to 100 pages"
    )

    if uploaded_file:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        try:
            # Text extraction based on file type
            if file_extension == ".pdf":
                extracted_text = extract_text_from_pdf(uploaded_file)
            elif file_extension == ".docx":
                extracted_text = extract_text_from_docx(uploaded_file)
            elif file_extension == ".txt":
                extracted_text = uploaded_file.getvalue().decode("utf-8")
            else:
                st.error("Unsupported file type.")
                return

            # Validate extracted text
            if len(extracted_text) < 100:
                st.warning("The extracted text is too short to summarize.")
                return

            # Display document metrics
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Total Words", len(extracted_text.split()))
            with col2:
                st.metric("Reading Time", f"{len(extracted_text.split()) // 200} mins")

            # Side-by-side document and summary
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("📄 Original Document")
                st.text_area("Full Text", extracted_text, height=400, disabled=True)
            
            with col2:
                st.subheader("🔍 Structured Summary")
                structured_summary = generate_structured_summary(extracted_text)
                st.markdown(structured_summary)

            # Download options
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "Download Original Text", 
                    extracted_text, 
                    "original_document.txt", 
                    "text/plain"
                )
            with col2:
                st.download_button(
                    "Download Summary", 
                    structured_summary, 
                    "document_summary.md", 
                    "text/markdown"
                )

        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

    else:
        st.info("👆 Upload a document to get started!")

# Run the app
if __name__ == "__main__":
    main()